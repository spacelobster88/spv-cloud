#!/usr/bin/env python3
"""
MIIT (工信部) vehicle announcement downloader and parser.

Downloads Word (.docx) documents from the MIIT Equipment Industry Development
Center (装备工业发展中心) announcement system, and parses vehicle records from
the structured tables within.

Usage:
    python miit_downloader.py --latest              # Download & parse latest batch
    python miit_downloader.py --batch 399           # Download & parse batch #399
    python miit_downloader.py --parse path/to.docx  # Parse an already-downloaded file
    python miit_downloader.py --list                # List available batches
"""

import argparse
import json
import logging
import os
import re
import sys
import time
import uuid
from pathlib import Path
from typing import Optional

import requests

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
)
logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------

# Base URL for the MIIT Equipment Industry Announcement Query System.
# The exact endpoints may change over time; update these as needed.
MIIT_BASE_URL = "https://ggzx.miit.gov.cn"

# Template for fetching the announcement list page.
# {page} is the page number (1-indexed).
BATCH_LIST_URL_TEMPLATE = f"{MIIT_BASE_URL}/GGzxFrontServlet?action=queryQCGGByPage&pageNo={{page}}"

# Template for downloading the announcement document for a given batch.
# {batch_number} is the integer batch number, e.g. 399.
BATCH_URL_TEMPLATE = f"{MIIT_BASE_URL}/GGzxFrontServlet?action=downLoadQCGG&batchNumber={{batch_number}}"

# Polite delay range (seconds) between HTTP requests.
REQUEST_DELAY_MIN = 2.0
REQUEST_DELAY_MAX = 3.0

# Default directory for downloaded documents (relative to project root).
DEFAULT_DATA_DIR = os.path.join(
    os.path.dirname(os.path.abspath(__file__)), "..", "data", "miit_docs"
)

# Default User-Agent header.
DEFAULT_USER_AGENT = (
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
    "AppleWebKit/537.36 (KHTML, like Gecko) "
    "Chrome/124.0.0.0 Safari/537.36"
)

# ---------------------------------------------------------------------------
# Chinese header -> English field mapping
# ---------------------------------------------------------------------------

HEADER_MAP = {
    "企业名称": "manufacturer",
    "产品商标": "brand",
    "产品型号": "model_number",
    "产品名称": "name",
    "外形尺寸(mm)": "overall_dimensions",
    "外形尺寸（mm）": "overall_dimensions",
    "货厢尺寸(mm)": "cargo_dimensions",
    "货厢尺寸（mm）": "cargo_dimensions",
    "总质量(kg)": "gvw",
    "总质量（kg）": "gvw",
    "整备质量(kg)": "curb_weight",
    "整备质量（kg）": "curb_weight",
    "额定载质量(kg)": "rated_payload",
    "额定载质量（kg）": "rated_payload",
    "底盘型号": "chassis_model",
    "底盘企业": "chassis_brand",
    "发动机型号": "engine_model",
    "发动机排量(ml)": "displacement_ml",
    "发动机排量（ml）": "displacement_ml",
    "发动机功率(kW)": "power_kw",
    "发动机功率（kW）": "power_kw",
    "排放标准": "emission_standard",
    "燃料种类": "fuel_type",
    "轴数": "axle_count",
    "轮胎规格": "tire_spec",
    "轮胎数": "tire_count",
    "驱动形式": "drive_type",
    "轴距(mm)": "wheelbase",
    "轴距（mm）": "wheelbase",
}


# ---------------------------------------------------------------------------
# MIITDownloader
# ---------------------------------------------------------------------------


class MIITDownloader:
    """Download MIIT vehicle announcement Word documents."""

    def __init__(self, data_dir: str = DEFAULT_DATA_DIR):
        self.data_dir = Path(data_dir)
        self.data_dir.mkdir(parents=True, exist_ok=True)
        self.session = requests.Session()
        self.session.headers.update({
            "User-Agent": DEFAULT_USER_AGENT,
            "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
            "Accept-Language": "zh-CN,zh;q=0.9,en;q=0.8",
            "Referer": MIIT_BASE_URL + "/",
        })

    # -- helpers --

    def _polite_delay(self):
        """Sleep for a polite interval between requests."""
        import random
        delay = random.uniform(REQUEST_DELAY_MIN, REQUEST_DELAY_MAX)
        logger.debug("Sleeping %.1fs between requests", delay)
        time.sleep(delay)

    def _download_file(self, url: str, dest: Path) -> Path:
        """Download a file from *url* and save it to *dest*."""
        logger.info("Downloading %s -> %s", url, dest)
        resp = self.session.get(url, timeout=60, stream=True)
        resp.raise_for_status()

        # Try to detect filename from Content-Disposition
        content_disp = resp.headers.get("Content-Disposition", "")
        if "filename" in content_disp:
            match = re.search(r'filename[*]?=["\']?([^"\';]+)', content_disp)
            if match:
                suggested = match.group(1).strip()
                logger.debug("Server suggested filename: %s", suggested)

        dest.parent.mkdir(parents=True, exist_ok=True)
        with open(dest, "wb") as fp:
            for chunk in resp.iter_content(chunk_size=8192):
                fp.write(chunk)

        file_size = dest.stat().st_size
        logger.info("Downloaded %s (%.1f KB)", dest.name, file_size / 1024)
        return dest

    # -- public API --

    def list_available_batches(self) -> list[dict]:
        """
        List available announcement batches from the MIIT website.

        Returns a list of dicts, each with at minimum:
            - batch_number (int)
            - title (str)
            - url (str)
        """
        batches: list[dict] = []
        try:
            from bs4 import BeautifulSoup
        except ImportError:
            logger.error(
                "beautifulsoup4 is required for listing batches. "
                "Install it with: pip install beautifulsoup4 lxml"
            )
            return batches

        # Fetch the first page of the announcement list.
        url = BATCH_LIST_URL_TEMPLATE.format(page=1)
        logger.info("Fetching batch list from %s", url)
        try:
            resp = self.session.get(url, timeout=30)
            resp.raise_for_status()
        except requests.RequestException as exc:
            logger.error("Failed to fetch batch list: %s", exc)
            return batches

        soup = BeautifulSoup(resp.text, "lxml")

        # The page typically lists batches as links. Try several selectors
        # since the page structure may vary.
        links = soup.select("a[href*='batch'], a[href*='QCGG'], a[href*='downLoad']")
        if not links:
            # Fallback: look for links whose text contains "批"
            links = [a for a in soup.find_all("a") if a.get_text() and "批" in a.get_text()]

        for link in links:
            text = link.get_text(strip=True)
            href = link.get("href", "")
            if not href.startswith("http"):
                href = MIIT_BASE_URL + ("/" if not href.startswith("/") else "") + href

            # Try to extract batch number from text like "第399批"
            m = re.search(r"第(\d+)批", text)
            batch_num = int(m.group(1)) if m else None

            batches.append({
                "batch_number": batch_num,
                "title": text,
                "url": href,
            })

        logger.info("Found %d batches", len(batches))
        return batches

    def download_batch(self, batch_number: int) -> Path:
        """
        Download the announcement document for a specific batch.

        Returns the Path to the saved .docx file.
        """
        dest = self.data_dir / f"batch_{batch_number}.docx"
        if dest.exists():
            logger.info("Batch %d already downloaded: %s", batch_number, dest)
            return dest

        url = BATCH_URL_TEMPLATE.format(batch_number=batch_number)
        self._polite_delay()
        return self._download_file(url, dest)

    def download_latest_batch(self) -> Path:
        """
        Discover and download the latest available announcement batch.

        Returns the Path to the saved .docx file.
        """
        batches = self.list_available_batches()
        if not batches:
            # If listing fails, try a heuristic: attempt recent batch numbers
            # starting from a reasonable recent number and working down.
            logger.warning(
                "Could not list batches from website; "
                "trying heuristic batch number detection"
            )
            return self._download_latest_heuristic()

        # Sort by batch_number descending, pick the highest.
        numbered = [b for b in batches if b["batch_number"] is not None]
        if not numbered:
            logger.error("No numbered batches found in listing")
            raise RuntimeError("Cannot determine latest batch number")

        numbered.sort(key=lambda b: b["batch_number"], reverse=True)
        latest = numbered[0]
        logger.info(
            "Latest batch: #%d (%s)", latest["batch_number"], latest["title"]
        )
        return self.download_batch(latest["batch_number"])

    def _download_latest_heuristic(self) -> Path:
        """
        Try to download the latest batch by probing recent batch numbers.

        Starts from batch 400 and works downward until a successful download.
        """
        for candidate in range(410, 380, -1):
            url = BATCH_URL_TEMPLATE.format(batch_number=candidate)
            try:
                logger.debug("Probing batch %d ...", candidate)
                resp = self.session.head(url, timeout=15, allow_redirects=True)
                if resp.status_code == 200:
                    return self.download_batch(candidate)
            except requests.RequestException:
                pass
            self._polite_delay()

        raise RuntimeError(
            "Could not discover latest batch via heuristic probing (tried 410-381)"
        )


# ---------------------------------------------------------------------------
# MIITDocParser
# ---------------------------------------------------------------------------


class MIITDocParser:
    """Parse a downloaded MIIT announcement .docx file into vehicle records."""

    def __init__(self, doc_path: Path, batch_number: Optional[int] = None):
        self.doc_path = Path(doc_path)
        # Try to infer batch number from filename if not supplied.
        if batch_number is None:
            m = re.search(r"batch_(\d+)", self.doc_path.stem)
            self.batch_number = int(m.group(1)) if m else None
        else:
            self.batch_number = batch_number

    # -- dimension / weight helpers --

    @staticmethod
    def _parse_dimensions(text: str) -> dict:
        """
        Parse a dimension string like '10500×2550×3990' or '10500*2550*3990'
        into length, width, height (as integers in mm).
        """
        if not text:
            return {}
        # Normalize separators
        normalized = re.sub(r"[×xX*\u00d7]", ",", text.strip())
        parts = [p.strip() for p in normalized.split(",") if p.strip()]
        nums: list[int] = []
        for p in parts:
            m = re.search(r"(\d+)", p)
            if m:
                nums.append(int(m.group(1)))
        if len(nums) >= 3:
            return {"l": nums[0], "w": nums[1], "h": nums[2]}
        return {}

    @staticmethod
    def _safe_int(text: str) -> Optional[int]:
        """Parse an integer from text, ignoring non-numeric characters."""
        if not text:
            return None
        m = re.search(r"(\d+)", text.strip())
        return int(m.group(1)) if m else None

    @staticmethod
    def _safe_float(text: str) -> Optional[float]:
        """Parse a float from text."""
        if not text:
            return None
        m = re.search(r"(\d+\.?\d*)", text.strip())
        return float(m.group(1)) if m else None

    # -- table parsing --

    def _extract_table_rows(self, table) -> list[list[str]]:
        """
        Extract all rows from a python-docx table as a list of string lists.

        Handles merged cells by repeating the merged cell value.
        """
        rows: list[list[str]] = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(cells)
        return rows

    def _table_to_kv_pairs(self, table) -> list[dict[str, str]]:
        """
        Convert a table whose structure is key-value pairs into records.

        MIIT announcement tables often use a two-column or multi-column
        layout where even columns are keys and odd columns are values,
        or where each row has a label and a value.

        This method handles both single-record (one vehicle per table) and
        multi-record (header row + data rows) formats.
        """
        rows = self._extract_table_rows(table)
        if not rows:
            return []

        # Heuristic: check if the first row looks like a header row
        # by checking whether multiple cells match known Chinese headers.
        first_row = rows[0]
        header_hits = sum(
            1 for cell in first_row
            if cell in HEADER_MAP or any(h in cell for h in HEADER_MAP)
        )

        if header_hits >= 3 and len(rows) > 1:
            # Multi-record table: first row is header, remaining are data.
            return self._parse_multi_record_table(rows)
        else:
            # Single-record table: key-value pairs laid out in rows.
            return [self._parse_single_record_table(rows)]

    def _parse_multi_record_table(self, rows: list[list[str]]) -> list[dict[str, str]]:
        """Parse a table where row 0 is headers and rows 1..N are records."""
        headers = rows[0]
        records: list[dict[str, str]] = []
        for row in rows[1:]:
            if not any(cell.strip() for cell in row):
                continue  # skip empty rows
            record: dict[str, str] = {}
            for idx, header in enumerate(headers):
                header_clean = header.strip()
                if idx < len(row) and header_clean:
                    record[header_clean] = row[idx].strip()
            if record:
                records.append(record)
        return records

    def _parse_single_record_table(self, rows: list[list[str]]) -> dict[str, str]:
        """
        Parse a table with key-value layout.

        Common patterns:
          - 2-column: [key, value]
          - 4-column: [key1, value1, key2, value2]
          - 6-column: [key1, value1, key2, value2, key3, value3]
        """
        record: dict[str, str] = {}
        for row in rows:
            # Process pairs of cells
            i = 0
            while i < len(row) - 1:
                key = row[i].strip()
                val = row[i + 1].strip()
                if key and key in HEADER_MAP:
                    record[key] = val
                    i += 2
                elif key:
                    # Check partial matches
                    matched = False
                    for known_header in HEADER_MAP:
                        if known_header in key:
                            record[known_header] = val
                            matched = True
                            break
                    i += 2 if matched or val else 1
                else:
                    i += 1
        return record

    def _raw_record_to_vehicle(self, raw: dict[str, str]) -> dict:
        """
        Convert a raw Chinese-keyed record dict into our vehicle schema.
        """
        vehicle: dict = {}

        # Map known headers
        for cn_key, en_key in HEADER_MAP.items():
            if cn_key in raw:
                vehicle[en_key] = raw[cn_key]

        # --- Post-processing ---

        # Parse overall dimensions
        dims = self._parse_dimensions(vehicle.pop("overall_dimensions", ""))
        vehicle["length"] = dims.get("l")
        vehicle["width"] = dims.get("w")
        vehicle["height"] = dims.get("h")

        # Parse cargo dimensions
        cargo_dims = self._parse_dimensions(vehicle.pop("cargo_dimensions", ""))
        vehicle["cargo_length"] = cargo_dims.get("l")
        vehicle["cargo_width"] = cargo_dims.get("w")
        vehicle["cargo_height"] = cargo_dims.get("h")

        # Numeric fields
        vehicle["gvw"] = self._safe_int(vehicle.get("gvw", ""))
        vehicle["curb_weight"] = self._safe_int(vehicle.get("curb_weight", ""))
        vehicle["rated_payload"] = self._safe_int(vehicle.get("rated_payload", ""))
        vehicle["power_kw"] = self._safe_int(vehicle.get("power_kw", ""))
        vehicle["axle_count"] = self._safe_int(vehicle.get("axle_count", ""))
        vehicle["tire_count"] = self._safe_int(vehicle.get("tire_count", ""))
        vehicle["wheelbase"] = self._safe_int(vehicle.get("wheelbase", ""))

        # Displacement: convert ml to liters
        disp_ml = self._safe_float(vehicle.pop("displacement_ml", ""))
        vehicle["displacement"] = round(disp_ml / 1000, 1) if disp_ml else None

        # Generate id
        vehicle["id"] = str(uuid.uuid4())

        # Announcement batch
        if self.batch_number is not None:
            vehicle["announcement_batch"] = f"第{self.batch_number}批"
        else:
            vehicle["announcement_batch"] = None

        # Defaults for fields not available in the doc
        vehicle.setdefault("images", [])
        vehicle.setdefault("is_tax_exempt", False)
        vehicle.setdefault("is_fuel_exempt", False)
        vehicle.setdefault("is_environmental", False)
        vehicle.setdefault("vehicle_type", None)
        vehicle.setdefault("vehicle_category", None)
        vehicle.setdefault("purpose", None)
        vehicle.setdefault("tonnage_class", None)
        vehicle.setdefault("engine_brand", None)
        vehicle.setdefault("power_hp", None)
        vehicle.setdefault("torque", None)
        vehicle.setdefault("cylinders", None)
        vehicle.setdefault("fuel_consumption", None)
        vehicle.setdefault("suspension_type", None)
        vehicle.setdefault("certificate_number", None)
        vehicle.setdefault("cargo_volume", None)
        vehicle.setdefault("payload", vehicle.get("rated_payload"))
        vehicle.setdefault("description", "")
        vehicle.setdefault("announcement_date", None)

        # Compute power_hp from power_kw if not set
        if vehicle.get("power_kw") and not vehicle.get("power_hp"):
            vehicle["power_hp"] = int(round(vehicle["power_kw"] * 1.3596))

        # Infer vehicle_type from name when possible
        if vehicle.get("name") and not vehicle.get("vehicle_type"):
            vehicle["vehicle_type"] = self._infer_vehicle_type(vehicle["name"])

        return vehicle

    @staticmethod
    def _infer_vehicle_type(name: str) -> Optional[str]:
        """Try to infer vehicle type from the product name."""
        type_keywords = [
            ("自卸车", "自卸"),
            ("搅拌车", "搅拌"),
            ("泵车", "泵车"),
            ("随车吊", "随车起重"),
            ("冷藏车", "冷藏"),
            ("油罐车", "油罐"),
            ("洒水车", "洒水"),
            ("垃圾车", "垃圾"),
            ("消防车", "消防"),
            ("清障车", "清障"),
            ("压缩垃圾车", "压缩式垃圾"),
            ("高空作业车", "高空作业"),
            ("厢式运输车", "厢式"),
            ("仓栅式运输车", "仓栅"),
            ("半挂牵引车", "牵引"),
            ("载货汽车", "载货"),
            ("平板运输车", "平板"),
        ]
        for vtype, keyword in type_keywords:
            if keyword in name:
                return vtype
        return None

    # -- main parse entry --

    def parse(self) -> list[dict]:
        """
        Parse the Word document and return a list of vehicle records.
        """
        try:
            from docx import Document
        except ImportError:
            logger.error(
                "python-docx is required for parsing. "
                "Install it with: pip install python-docx"
            )
            return []

        if not self.doc_path.exists():
            logger.error("Document not found: %s", self.doc_path)
            return []

        logger.info("Parsing document: %s", self.doc_path)
        doc = Document(str(self.doc_path))

        vehicles: list[dict] = []
        table_count = len(doc.tables)
        logger.info("Found %d tables in document", table_count)

        for idx, table in enumerate(doc.tables):
            try:
                raw_records = self._table_to_kv_pairs(table)
                for raw in raw_records:
                    if not raw:
                        continue
                    vehicle = self._raw_record_to_vehicle(raw)
                    # Only include records that have at least a model number or name
                    if vehicle.get("model_number") or vehicle.get("name"):
                        vehicles.append(vehicle)
            except Exception as exc:
                logger.warning(
                    "Error parsing table %d/%d: %s", idx + 1, table_count, exc
                )
                continue

        logger.info(
            "Parsed %d vehicle records from %d tables", len(vehicles), table_count
        )
        return vehicles

    def _parse_new_vehicles_table(self, table) -> list[dict]:
        """
        Parse a table section specifically for new vehicle entries (新产品公告).

        This is a convenience wrapper around the generic table parser for
        tables that are known to contain new product announcements.
        """
        raw_records = self._table_to_kv_pairs(table)
        vehicles = []
        for raw in raw_records:
            if not raw:
                continue
            vehicle = self._raw_record_to_vehicle(raw)
            vehicle["announcement_type"] = "新产品公告"
            if vehicle.get("model_number") or vehicle.get("name"):
                vehicles.append(vehicle)
        return vehicles

    def _parse_change_vehicles_table(self, table) -> list[dict]:
        """
        Parse a table section for change/extension entries (变更扩展).

        This is a convenience wrapper around the generic table parser for
        tables that contain change or extension announcements.
        """
        raw_records = self._table_to_kv_pairs(table)
        vehicles = []
        for raw in raw_records:
            if not raw:
                continue
            vehicle = self._raw_record_to_vehicle(raw)
            vehicle["announcement_type"] = "变更扩展"
            if vehicle.get("model_number") or vehicle.get("name"):
                vehicles.append(vehicle)
        return vehicles


# ---------------------------------------------------------------------------
# CLI entry point
# ---------------------------------------------------------------------------


def main():
    parser = argparse.ArgumentParser(
        description="Download and parse MIIT vehicle announcement documents."
    )
    group = parser.add_mutually_exclusive_group(required=True)
    group.add_argument(
        "--latest",
        action="store_true",
        help="Download and parse the latest announcement batch.",
    )
    group.add_argument(
        "--batch",
        type=int,
        metavar="NUMBER",
        help="Download and parse a specific batch by number.",
    )
    group.add_argument(
        "--parse",
        type=str,
        metavar="FILE",
        help="Parse an already-downloaded .docx file.",
    )
    group.add_argument(
        "--list",
        action="store_true",
        help="List available announcement batches.",
    )
    parser.add_argument(
        "--data-dir",
        type=str,
        default=DEFAULT_DATA_DIR,
        help="Directory for downloaded documents (default: %(default)s).",
    )
    parser.add_argument(
        "--output",
        type=str,
        default=None,
        metavar="FILE",
        help="Write parsed vehicle records to a JSON file.",
    )
    parser.add_argument(
        "--verbose", "-v",
        action="store_true",
        help="Enable debug logging.",
    )

    args = parser.parse_args()

    if args.verbose:
        logging.getLogger().setLevel(logging.DEBUG)

    downloader = MIITDownloader(data_dir=args.data_dir)

    if args.list:
        batches = downloader.list_available_batches()
        if not batches:
            logger.warning("No batches found.")
            sys.exit(1)
        for b in batches:
            num = b.get("batch_number", "?")
            print(f"  Batch #{num}: {b['title']}")
        return

    # Determine the document path.
    doc_path: Optional[Path] = None

    if args.latest:
        try:
            doc_path = downloader.download_latest_batch()
        except Exception as exc:
            logger.error("Failed to download latest batch: %s", exc)
            sys.exit(1)

    elif args.batch:
        try:
            doc_path = downloader.download_batch(args.batch)
        except Exception as exc:
            logger.error("Failed to download batch %d: %s", args.batch, exc)
            sys.exit(1)

    elif args.parse:
        doc_path = Path(args.parse)
        if not doc_path.exists():
            logger.error("File not found: %s", doc_path)
            sys.exit(1)

    # Parse the document.
    if doc_path is None:
        logger.error("No document to parse.")
        sys.exit(1)

    # Determine batch number for the parser.
    batch_number = args.batch if args.batch else None
    doc_parser = MIITDocParser(doc_path, batch_number=batch_number)
    vehicles = doc_parser.parse()

    if not vehicles:
        logger.warning("No vehicle records parsed from %s", doc_path)
        sys.exit(0)

    # Summary
    brands = set(v.get("brand") for v in vehicles if v.get("brand"))
    types = set(v.get("vehicle_type") for v in vehicles if v.get("vehicle_type"))
    logger.info(
        "Parsed %d vehicles, %d brands, %d vehicle types",
        len(vehicles), len(brands), len(types),
    )

    # Output
    if args.output:
        out_path = Path(args.output)
        out_path.parent.mkdir(parents=True, exist_ok=True)
        with open(out_path, "w", encoding="utf-8") as fp:
            json.dump(vehicles, fp, ensure_ascii=False, indent=2)
        logger.info("Wrote %d records to %s", len(vehicles), out_path)
    else:
        # Print to stdout
        print(json.dumps(vehicles, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
