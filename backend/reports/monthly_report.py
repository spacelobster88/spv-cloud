"""Monthly report generator — produces Word documents with charts and tables.

Usage (CLI):
    python -m reports.monthly_report --year 2026 --month 3

The generated .docx file is saved to ``backend/reports/output/``.
"""

import argparse
import json
import logging
import tempfile
from collections import Counter
from datetime import datetime
from pathlib import Path
from typing import Any

import matplotlib
matplotlib.use("Agg")  # non-interactive backend
import matplotlib.pyplot as plt
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Project color palette
# ---------------------------------------------------------------------------
COLOR_PRIMARY = "#2D8CA0"
COLOR_SECONDARY = "#3AAFC9"
COLOR_DARK = "#1B2B4B"
COLOR_LIGHT = "#E8F4F8"
COLOR_ACCENT = "#F0A830"

_PALETTE = [
    COLOR_PRIMARY, COLOR_SECONDARY, COLOR_DARK,
    "#4DC9E6", "#F0A830", "#E8684A", "#6BBF8A",
    "#9B8EC4", "#D4A574", "#7ECFC0", "#C97B84",
]

# ---------------------------------------------------------------------------
# Certificate-number prefix -> province mapping
# ---------------------------------------------------------------------------
_CERT_PROVINCE_MAP: dict[str, str] = {
    "京": "北京", "津": "天津", "沪": "上海", "渝": "重庆",
    "冀": "河北", "豫": "河南", "云": "云南", "辽": "辽宁",
    "黑": "黑龙江", "湘": "湖南", "皖": "安徽", "鲁": "山东",
    "新": "新疆", "苏": "江苏", "浙": "浙江", "赣": "江西",
    "鄂": "湖北", "桂": "广西", "甘": "甘肃", "晋": "山西",
    "蒙": "内蒙古", "陕": "陕西", "吉": "吉林", "闽": "福建",
    "贵": "贵州", "粤": "广东", "川": "四川", "青": "青海",
    "藏": "西藏", "琼": "海南", "宁": "宁夏",
}


# ---------------------------------------------------------------------------
# Matplotlib Chinese font setup
# ---------------------------------------------------------------------------

def _setup_chinese_font() -> None:
    """Try to configure matplotlib for Chinese text rendering."""
    import matplotlib.font_manager as fm

    # Preferred fonts in order
    candidates = [
        "SimHei", "Heiti SC", "Heiti TC", "PingFang SC",
        "Microsoft YaHei", "WenQuanYi Micro Hei", "Noto Sans CJK SC",
        "STHeiti", "Arial Unicode MS",
    ]
    available = {f.name for f in fm.fontManager.ttflist}
    chosen = None
    for name in candidates:
        if name in available:
            chosen = name
            break

    if chosen:
        plt.rcParams["font.sans-serif"] = [chosen] + plt.rcParams["font.sans-serif"]
        logger.info("Using Chinese font: %s", chosen)
    else:
        logger.warning(
            "No Chinese font found — chart labels may render as boxes. "
            "Install SimHei or Noto Sans CJK SC for proper rendering."
        )
    plt.rcParams["axes.unicode_minus"] = False


# ---------------------------------------------------------------------------
# MonthlyReportGenerator
# ---------------------------------------------------------------------------

class MonthlyReportGenerator:
    """Generates a monthly statistical analysis report as a Word document."""

    def __init__(
        self,
        data_path: str = "backend/data/vehicles.json",
        output_dir: str = "backend/reports/output",
    ) -> None:
        """Initialize with data source and output directory."""
        self.data_path = Path(data_path)
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)

    # ------------------------------------------------------------------
    # Public API
    # ------------------------------------------------------------------

    def generate(self, year: int, month: int) -> Path:
        """Generate a complete monthly report as Word document.

        Returns the path to the generated ``.docx`` file.
        """
        logger.info("Generating monthly report for %04d-%02d …", year, month)

        vehicles = self._load_data()
        filtered = self._filter_by_month(vehicles, year, month)

        if not filtered:
            logger.warning("No vehicle data for %04d-%02d.", year, month)

        stats: dict[str, dict[str, int]] = {
            "region": self._stats_by_region(filtered),
            "usage": self._stats_by_usage(filtered),
            "type": self._stats_by_type(filtered),
            "brand": self._stats_by_brand(filtered),
            "emission": self._stats_by_emission(filtered),
            "fuel": self._stats_by_fuel(filtered),
        }

        _setup_chinese_font()
        charts = self._generate_charts(stats)

        output_path = self._build_word_doc(year, month, stats, charts, filtered)
        logger.info("Report saved to %s", output_path)
        return output_path

    # ------------------------------------------------------------------
    # Data loading & filtering
    # ------------------------------------------------------------------

    def _load_data(self) -> list[dict[str, Any]]:
        """Load vehicle data from JSON file."""
        logger.info("Loading data from %s", self.data_path)
        if not self.data_path.exists():
            raise FileNotFoundError(f"Data file not found: {self.data_path}")
        with open(self.data_path, encoding="utf-8") as fh:
            data: list[dict[str, Any]] = json.load(fh)
        logger.info("Loaded %d vehicle records.", len(data))
        return data

    def _filter_by_month(
        self, vehicles: list[dict[str, Any]], year: int, month: int
    ) -> list[dict[str, Any]]:
        """Filter vehicles whose *announcement_date* falls in the given month."""
        result: list[dict[str, Any]] = []
        prefix = f"{year:04d}-{month:02d}"
        for v in vehicles:
            date_str = v.get("announcement_date", "")
            if date_str.startswith(prefix):
                result.append(v)
        logger.info(
            "Filtered %d vehicles for %s (from %d total).",
            len(result), prefix, len(vehicles),
        )
        return result

    # ------------------------------------------------------------------
    # Statistics helpers
    # ------------------------------------------------------------------

    @staticmethod
    def _province_from_cert(cert: str) -> str:
        """Extract province name from certificate_number prefix."""
        if cert:
            prefix_char = cert[0]
            return _CERT_PROVINCE_MAP.get(prefix_char, "其他")
        return "未知"

    def _stats_by_region(self, vehicles: list[dict[str, Any]]) -> dict[str, int]:
        """Statistics grouped by province/region (derived from certificate)."""
        counter: Counter[str] = Counter()
        for v in vehicles:
            province = self._province_from_cert(v.get("certificate_number", ""))
            counter[province] += 1
        return dict(counter.most_common())

    def _stats_by_usage(self, vehicles: list[dict[str, Any]]) -> dict[str, int]:
        """Statistics grouped by vehicle usage category (purpose)."""
        counter: Counter[str] = Counter()
        for v in vehicles:
            counter[v.get("purpose", "未知")] += 1
        return dict(counter.most_common())

    def _stats_by_type(self, vehicles: list[dict[str, Any]]) -> dict[str, int]:
        """Statistics grouped by vehicle_type."""
        counter: Counter[str] = Counter()
        for v in vehicles:
            counter[v.get("vehicle_type", "未知")] += 1
        return dict(counter.most_common())

    def _stats_by_brand(self, vehicles: list[dict[str, Any]]) -> dict[str, int]:
        """Statistics grouped by brand."""
        counter: Counter[str] = Counter()
        for v in vehicles:
            counter[v.get("brand", "未知")] += 1
        return dict(counter.most_common())

    def _stats_by_emission(self, vehicles: list[dict[str, Any]]) -> dict[str, int]:
        """Statistics grouped by emission_standard."""
        counter: Counter[str] = Counter()
        for v in vehicles:
            counter[v.get("emission_standard", "未知")] += 1
        return dict(counter.most_common())

    def _stats_by_fuel(self, vehicles: list[dict[str, Any]]) -> dict[str, int]:
        """Statistics grouped by fuel_type."""
        counter: Counter[str] = Counter()
        for v in vehicles:
            counter[v.get("fuel_type", "未知")] += 1
        return dict(counter.most_common())

    # ------------------------------------------------------------------
    # Chart generation
    # ------------------------------------------------------------------

    def _generate_charts(self, stats: dict[str, dict[str, int]]) -> dict[str, Path]:
        """Generate chart images using matplotlib for each stat category."""
        tmp_dir = Path(tempfile.mkdtemp(prefix="spv_charts_"))
        charts: dict[str, Path] = {}

        chart_configs: list[dict[str, Any]] = [
            {
                "key": "region",
                "title": "按区域分布",
                "kind": "bar",
            },
            {
                "key": "usage",
                "title": "按用途类别分布",
                "kind": "pie",
            },
            {
                "key": "type",
                "title": "按车型分布",
                "kind": "bar",
            },
            {
                "key": "brand",
                "title": "按品牌分布",
                "kind": "barh",
            },
            {
                "key": "emission",
                "title": "按排放标准分布",
                "kind": "pie",
            },
            {
                "key": "fuel",
                "title": "按燃料类型分布",
                "kind": "pie",
            },
        ]

        for cfg in chart_configs:
            key = cfg["key"]
            data = stats.get(key, {})
            if not data:
                logger.info("Skipping chart '%s' — no data.", key)
                continue

            path = tmp_dir / f"{key}.png"
            self._render_chart(
                data=data,
                title=cfg["title"],
                kind=cfg["kind"],
                output_path=path,
            )
            charts[key] = path

        return charts

    def _render_chart(
        self,
        data: dict[str, int],
        title: str,
        kind: str,
        output_path: Path,
    ) -> None:
        """Render a single chart and save to *output_path*."""
        labels = list(data.keys())
        values = list(data.values())
        colors = [_PALETTE[i % len(_PALETTE)] for i in range(len(labels))]

        fig, ax = plt.subplots(figsize=(8, 5))

        if kind == "bar":
            bars = ax.bar(labels, values, color=colors, edgecolor="white")
            ax.set_ylabel("数量", fontsize=11)
            # Rotate long labels
            if any(len(l) > 3 for l in labels):
                plt.xticks(rotation=35, ha="right", fontsize=9)
            # Value labels on bars
            for bar, val in zip(bars, values):
                ax.text(
                    bar.get_x() + bar.get_width() / 2, bar.get_height() + 0.3,
                    str(val), ha="center", va="bottom", fontsize=9,
                )

        elif kind == "barh":
            # Show top-15 for horizontal bar
            if len(labels) > 15:
                labels = labels[:15]
                values = values[:15]
                colors = colors[:15]
            labels = labels[::-1]
            values = values[::-1]
            colors = colors[::-1]
            bars = ax.barh(labels, values, color=colors, edgecolor="white")
            ax.set_xlabel("数量", fontsize=11)
            for bar, val in zip(bars, values):
                ax.text(
                    bar.get_width() + 0.2, bar.get_y() + bar.get_height() / 2,
                    str(val), ha="left", va="center", fontsize=9,
                )

        elif kind == "pie":
            wedges, texts, autotexts = ax.pie(
                values, labels=labels, autopct="%1.1f%%",
                colors=colors, startangle=140, pctdistance=0.8,
            )
            for t in autotexts:
                t.set_fontsize(9)

        ax.set_title(title, fontsize=14, fontweight="bold", color=COLOR_DARK)
        fig.tight_layout()
        fig.savefig(output_path, dpi=150, bbox_inches="tight")
        plt.close(fig)
        logger.info("Chart saved: %s", output_path)

    # ------------------------------------------------------------------
    # Word document generation
    # ------------------------------------------------------------------

    def _build_word_doc(
        self,
        year: int,
        month: int,
        stats: dict[str, dict[str, int]],
        charts: dict[str, Path],
        vehicles: list[dict[str, Any]],
    ) -> Path:
        """Build the Word document with python-docx."""
        doc = Document()

        # -- Page setup --------------------------------------------------
        section = doc.sections[0]
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)

        # -- Header / footer --------------------------------------------
        header_para = section.header.paragraphs[0]
        header_para.text = f"专用汽车公告产品月度分析报告 — {year}年{month}月"
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._set_run_font(header_para.runs[0], size=Pt(8), color=RGBColor(0x99, 0x99, 0x99))

        footer_para = section.footer.paragraphs[0]
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._add_page_number(footer_para)

        # -- Title page --------------------------------------------------
        self._add_empty_paragraphs(doc, 6)
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("专用汽车公告产品\n月度分析报告")
        self._set_run_font(run, size=Pt(28), bold=True, color=RGBColor(0x1B, 0x2B, 0x4B))

        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run(f"{year}年{month}月")
        self._set_run_font(run, size=Pt(18), color=RGBColor(0x2D, 0x8C, 0xA0))

        generated_para = doc.add_paragraph()
        generated_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = generated_para.add_run(f"生成日期：{datetime.now().strftime('%Y-%m-%d')}")
        self._set_run_font(run, size=Pt(10), color=RGBColor(0x99, 0x99, 0x99))

        doc.add_page_break()

        # -- Executive summary -------------------------------------------
        total = len(vehicles)
        doc.add_heading("概要", level=1)
        summary_lines = [
            f"本报告统计了{year}年{month}月公告发布的专用汽车产品信息。",
            f"本月新增公告产品共计 {total} 款。",
        ]
        if stats["brand"]:
            top_brands = list(stats["brand"].items())[:5]
            brand_text = "、".join(f"{b}({c}款)" for b, c in top_brands)
            summary_lines.append(f"品牌分布方面，排名前五的品牌为：{brand_text}。")
        if stats["type"]:
            top_types = list(stats["type"].items())[:3]
            type_text = "、".join(f"{t}({c}款)" for t, c in top_types)
            summary_lines.append(f"车型方面，主要类型为：{type_text}。")
        if stats["emission"]:
            top_emission = list(stats["emission"].items())[0]
            summary_lines.append(
                f"排放标准以{top_emission[0]}为主，占比"
                f"{top_emission[1] / total * 100:.1f}%。" if total else ""
            )
        for line in summary_lines:
            if line:
                doc.add_paragraph(line)

        doc.add_page_break()

        # -- Section 1: By region ----------------------------------------
        doc.add_heading("第一章  按区域分析", level=1)
        doc.add_paragraph("以下为本月新增公告产品的区域（省份）分布情况。区域信息根据产品合格证号前缀推导。")
        self._add_stats_table(doc, stats["region"], "省份", "数量")
        if "region" in charts:
            doc.add_picture(str(charts["region"]), width=Inches(5.5))
            self._center_last_paragraph(doc)

        doc.add_page_break()

        # -- Section 2: By usage -----------------------------------------
        doc.add_heading("第二章  按用途类别分析", level=1)
        doc.add_paragraph("以下为按车辆用途分类的统计结果。")
        self._add_stats_table(doc, stats["usage"], "用途类别", "数量")
        if "usage" in charts:
            doc.add_picture(str(charts["usage"]), width=Inches(4.5))
            self._center_last_paragraph(doc)

        doc.add_page_break()

        # -- Section 3: By type ------------------------------------------
        doc.add_heading("第三章  按车型分析", level=1)
        doc.add_paragraph("以下为按车辆类型的统计结果。")
        self._add_stats_table(doc, stats["type"], "车辆类型", "数量")
        if "type" in charts:
            doc.add_picture(str(charts["type"]), width=Inches(5.5))
            self._center_last_paragraph(doc)

        doc.add_page_break()

        # -- Section 4: By brand -----------------------------------------
        doc.add_heading("第四章  按品牌分析", level=1)
        doc.add_paragraph("以下为按品牌的市场份额统计。")
        self._add_stats_table(doc, stats["brand"], "品牌", "数量")
        if "brand" in charts:
            doc.add_picture(str(charts["brand"]), width=Inches(5.5))
            self._center_last_paragraph(doc)

        doc.add_page_break()

        # -- Section 5: Emission standard --------------------------------
        doc.add_heading("第五章  排放标准分析", level=1)
        doc.add_paragraph("以下为本月新增产品的排放标准分布情况。")
        self._add_stats_table(doc, stats["emission"], "排放标准", "数量")
        if "emission" in charts:
            doc.add_picture(str(charts["emission"]), width=Inches(4.5))
            self._center_last_paragraph(doc)

        doc.add_page_break()

        # -- Section 6: Fuel type ----------------------------------------
        doc.add_heading("第六章  燃料类型分析", level=1)
        doc.add_paragraph("以下为本月新增产品的燃料类型分布情况。")
        self._add_stats_table(doc, stats["fuel"], "燃料类型", "数量")
        if "fuel" in charts:
            doc.add_picture(str(charts["fuel"]), width=Inches(4.5))
            self._center_last_paragraph(doc)

        doc.add_page_break()

        # -- Appendix: Full vehicle list ---------------------------------
        doc.add_heading("附录  本月新增公告产品清单", level=1)
        if vehicles:
            self._add_vehicle_list_table(doc, vehicles)
        else:
            doc.add_paragraph("本月无新增公告产品。")

        # -- Save --------------------------------------------------------
        filename = f"monthly_report_{year:04d}_{month:02d}.docx"
        output_path = self.output_dir / filename
        doc.save(str(output_path))
        return output_path

    # ------------------------------------------------------------------
    # Word-doc helper methods
    # ------------------------------------------------------------------

    @staticmethod
    def _set_run_font(
        run: Any,
        size: Pt | None = None,
        bold: bool = False,
        color: RGBColor | None = None,
    ) -> None:
        """Apply font settings to a run."""
        run.bold = bold
        if size:
            run.font.size = size
        if color:
            run.font.color.rgb = color

    @staticmethod
    def _add_empty_paragraphs(doc: Document, count: int) -> None:
        for _ in range(count):
            doc.add_paragraph("")

    @staticmethod
    def _center_last_paragraph(doc: Document) -> None:
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    @staticmethod
    def _add_page_number(paragraph: Any) -> None:
        """Insert an auto page-number field into a paragraph."""
        run = paragraph.add_run()
        fldChar1 = run._r.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "begin"})
        run._r.append(fldChar1)

        run2 = paragraph.add_run()
        instrText = run2._r.makeelement(qn("w:instrText"), {})
        instrText.text = " PAGE "
        run2._r.append(instrText)

        run3 = paragraph.add_run()
        fldChar2 = run3._r.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "end"})
        run3._r.append(fldChar2)

    def _add_stats_table(
        self, doc: Document, data: dict[str, int], col1_title: str, col2_title: str
    ) -> None:
        """Add a two-column stats table with alternating row shading."""
        if not data:
            doc.add_paragraph("暂无数据。")
            return

        total = sum(data.values())
        table = doc.add_table(rows=1, cols=3)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"

        # Header row
        hdr = table.rows[0]
        for idx, text in enumerate([col1_title, col2_title, "占比"]):
            cell = hdr.cells[idx]
            cell.text = text
            self._shade_cell(cell, "1B2B4B")
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    r.bold = True
                    r.font.size = Pt(10)

        # Data rows
        for i, (label, count) in enumerate(data.items()):
            row = table.add_row()
            row.cells[0].text = label
            row.cells[1].text = str(count)
            pct = f"{count / total * 100:.1f}%" if total else "0%"
            row.cells[2].text = pct

            # Alternating shading
            if i % 2 == 0:
                for cell in row.cells:
                    self._shade_cell(cell, "E8F4F8")

            # Center-align number columns
            for ci in (1, 2):
                for p in row.cells[ci].paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Total row
        total_row = table.add_row()
        total_row.cells[0].text = "合计"
        total_row.cells[1].text = str(total)
        total_row.cells[2].text = "100%"
        for cell in total_row.cells:
            self._shade_cell(cell, "2D8CA0")
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    r.bold = True

    def _add_vehicle_list_table(
        self, doc: Document, vehicles: list[dict[str, Any]]
    ) -> None:
        """Add an appendix table listing all vehicles with key fields."""
        columns = [
            ("序号", ""),
            ("产品名称", "name"),
            ("品牌", "brand"),
            ("型号", "model_number"),
            ("车辆类型", "vehicle_type"),
            ("用途", "purpose"),
            ("排放标准", "emission_standard"),
            ("燃料类型", "fuel_type"),
            ("公告日期", "announcement_date"),
        ]

        table = doc.add_table(rows=1, cols=len(columns))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"

        # Header
        hdr = table.rows[0]
        for idx, (title, _) in enumerate(columns):
            cell = hdr.cells[idx]
            cell.text = title
            self._shade_cell(cell, "1B2B4B")
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    r.bold = True
                    r.font.size = Pt(8)

        # Data
        for i, v in enumerate(vehicles):
            row = table.add_row()
            for col_idx, (_, field) in enumerate(columns):
                if col_idx == 0:
                    row.cells[col_idx].text = str(i + 1)
                else:
                    row.cells[col_idx].text = str(v.get(field, ""))
                for p in row.cells[col_idx].paragraphs:
                    for r in p.runs:
                        r.font.size = Pt(8)

            if i % 2 == 0:
                for cell in row.cells:
                    self._shade_cell(cell, "E8F4F8")

    @staticmethod
    def _shade_cell(cell: Any, hex_color: str) -> None:
        """Apply background shading to a table cell."""
        shading = cell._tc.get_or_add_tcPr().makeelement(
            qn("w:shd"),
            {
                qn("w:val"): "clear",
                qn("w:color"): "auto",
                qn("w:fill"): hex_color,
            },
        )
        cell._tc.get_or_add_tcPr().append(shading)


# ---------------------------------------------------------------------------
# CLI entry point
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    logging.basicConfig(
        level=logging.INFO,
        format="%(asctime)s [%(levelname)s] %(name)s: %(message)s",
    )

    parser = argparse.ArgumentParser(
        description="Generate a monthly SPV report (.docx).",
    )
    parser.add_argument("--year", type=int, required=True, help="Report year, e.g. 2026")
    parser.add_argument("--month", type=int, required=True, help="Report month (1-12)")
    parser.add_argument(
        "--data", type=str, default="backend/data/vehicles.json",
        help="Path to vehicles.json",
    )
    parser.add_argument(
        "--output", type=str, default="backend/reports/output",
        help="Output directory for the generated report",
    )
    args = parser.parse_args()

    generator = MonthlyReportGenerator(data_path=args.data, output_dir=args.output)
    result = generator.generate(year=args.year, month=args.month)
    print(f"Report generated: {result}")
